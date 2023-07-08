!pip install python-docx
import docx
import pandas as pd
import numpy as np
import locale
import csv
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
locale.setlocale(locale.LC_ALL, 'en_US.UTF-8')  # Replace 'en_US.UTF-8' with your desired locale

#file_name = ["insert list of accounts here"]
#data_combined = "/content/data_combined.csv"

c_df = pd.read_csv(data_combined, header=None)

headers = ['account', 'ac_code', 'desc', 'date', 'principle', 'interest', 'misc',
           'amount', 'balance','proceeds','sale fee', 'key fees', 'transport fee', 'auction other fee',
           'auction repair fees', 'recondition fee','overpayment','refund', 'full account', 'note date',
           'sold date', 'name', 'address', 'city', 'state', 'zip', 'year', 'make', 'model', 'vin']

c_df.columns = headers

csv_filename = str(file_name) + ".csv"
csv_path = '/content/' + csv_filename

grouped = c_df.groupby('account')

# Create a dictionary to store DataFrames for each account
account_dfs = {}

# Iterate over each account group and store it in the dictionary
for account_number, group in grouped:
    # Store the group DataFrame in the dictionary
    account_dfs[account_number] = group

account = account_dfs[file_name]
df = account

docx_filename = "defi_letter.docx"
docx_path = '/content/' + docx_filename

newcsv_filename = str(file_name) + ".docx"
newdocx_path = '/content/updated/' + newcsv_filename

doc = docx.Document(docx_path)

if 'UACG' in df['desc'].values:
  bal = df.loc[df['desc'] == 'UACG', 'balance'].values.item()
  uacg = df.loc[df['desc'] == 'UACG', 'amount'].values.item()*-1
else:
  bal =  df.loc[(df['ac_code'] == '30-C') & (df['desc'] == 'PYSS'), 'balance'].values.item()
  uacg = 0

payoff_date_test = df.loc[(df['ac_code'] == '30-C') & (df['desc'] == 'PYSS'), 'date'].values.item()
condition = df['date'] == payoff_date_test


int_part1 = df.loc[(df['ac_code'] == '30-C') & (df['desc'] == 'PYSS'), 'interest'].values
int_part2 = df.loc[(df['ac_code'] == '30-R') & (df['desc'] == 'PYSS') , 'interest'].values
misc = df.loc[(df['ac_code'] == '30-C') & (df['desc'] == 'PYSS'), 'misc'].values
repo = df.loc[(df['desc'] == 'REPO') & condition, 'misc'].values
impo = df.loc[(df['desc'] == 'IMPO') & condition, 'misc'].values
auct = df.loc[(df['desc'] == 'AUCT') & condition, 'misc'].values
nsf = df.loc[(df['desc'] == '0NF2') & condition, 'misc'].values

list = [int_part1, int_part2, misc, repo, impo, auct, nsf]
new_list = []

for i in range(len(list)):
  value = list[i]
  if len(value) > 0:
    result = value[0]
    numeric_value = result.item()
    new_list.append(value[0].item())
  else:
    result = 0
    new_list.append(result)

int_part1 = new_list[0]
int_part2 = new_list[1]
misc = new_list[2]
repo = new_list[3]
impo = new_list[4]
auct = new_list[5]
nsf = new_list[6]

interest = int_part1 + int_part2
other_cost = repo + impo
other_other_cost = auct + nsf
net_bal = round(float(bal + interest + misc), 2)

proceeds = df.iloc[1,9]
sale_fee = df.iloc[1,10]
key_fee = df.iloc[1,11]
trans_fee = df.iloc[1,12]
auct_oth_fee = df.iloc[1,13]
auct_rep_fee = df.iloc[1,14]
recon_fee = df.iloc[1,15]

subtotal = round((net_bal - proceeds), 2)
storing = 0
cos_veh = sale_fee
cop_veh = trans_fee + recon_fee + auct_oth_fee + key_fee  + auct_rep_fee
court_cost = 0
total_cost = round((cos_veh + cop_veh + other_cost + other_other_cost), 2)
balance_due = round(float(subtotal + total_cost + uacg), 2)

if df.iloc[1,16] is None:
  refund =  df.iloc[1,17]
else:
  refund = df.iloc[1,16]

if abs(balance_due) != refund:
  balance_due += 5
  cop_veh += 5
  total_cost += 5
  if abs(balance_due) != refund:
    balance_due += 25
    cop_veh += 25
    total_cost += 25
    if abs(balance_due) != refund:
      pass
    else:
      pass
  else:
    pass
else:
  pass

table1 = doc.tables[1]

prin_cell = table1.cell(0,3)
int_cell = table1.cell(1,3)
late_fee_cell = table1.cell(2,3)
net_bal_cell = table1.cell(3,4)

proceeds_cell = table1.cell(4,3)

subtotal_cell = table1.cell(5,4)

repo_cell = table1.cell(6,3)
impound_cell = table1.cell(7,3)
storing_cell = table1.cell(8,3)
preparing_cell = table1.cell(9,3)
selling_cell  = table1.cell(10,3)
court_cost_cell = table1.cell(11,3)
other_cost_cell = table1.cell(12,3)
total_cost_cell = table1.cell(13,4)

gap_cell = table1.cell(15,3)

total_credit_cell = table1.cell(18,4)

surplus_cell = table1.cell(19,4)

account_number =  str(df.iloc[1,18])

formatted_account_number = (
    account_number[:3] + "-" +
    account_number[3:6] + "-" +
    account_number[6:13] + "-" +
    account_number[13:]
)

today = datetime.today().strftime("%m/%d/%Y")

note_date = df.iloc[1,19]
# date_obj1 = datetime.strptime(note_date, "%m-%d-%Y")
date_obj1 = datetime.strptime(note_date, "%Y-%m-%d")
formatted_note_date = date_obj1.strftime("%m/%d/%Y")

sold_date = df.iloc[1,20]
date_obj2 = datetime.strptime(sold_date, "%Y-%m-%d")
formatted_sold_date = date_obj2.strftime("%m/%d/%Y")

payoff_date = df.loc[df['ac_code'] == '30-C', 'date']
payoff_date_value = payoff_date.to_string(index=False)
date_obj = datetime.strptime(payoff_date_value, "%Y-%m-%d")
formatted_payoff_date = date_obj.strftime("%m/%d/%Y")

name = df.iloc[1,21]
address = df.iloc[1,22]
city = df.iloc[1,23]
state = df.iloc[1,24]
zip = df.iloc[1,25]
year = df.iloc[1,26]
make = df.iloc[1,27]
model = df.iloc[1,28]
VIN = df.iloc[1,29]

def insert_text(paragraph_index, runs_index, text):
  paragraph = doc.paragraphs[paragraph_index]
  paragraph.text = text
  # paragraph.runs[runs_index].font.bold = True
  paragraph.runs[runs_index].font.size = Pt(10)

insert_text(2, 0, "Date: " + today)
insert_text(8, 0, name)
insert_text(9, 0, str(address))
insert_text(10, 0, str(city) + ", " + str(state) + " " +  str(zip))
insert_text(13, 0, 'Re:      Account No. ' + str(formatted_account_number))
insert_text(14, 0, '            Retail Installment Sale or Credit Sale Contract dated ' +  str(formatted_note_date) + ' ("Agreement")')
insert_text(15, 0, '            ' + str(year) + ' ' + str(make) + ' ' + str(model) +  ', ' + str(VIN) + ' ("Vehicle")')
insert_text(17, 0, 'Dear ' + name)
insert_text(19, 0, 'Please be advised that we disposed of the Vehicle on ' + str(formatted_sold_date) + '.  The proceeds of the sale have been applied as explained below....')

payoff_date_cell = doc.tables[1].rows[0].cells[1]
payoff_date_cell.text = 'Unpaid Principal Balance of Agreement as of ' + str(formatted_payoff_date)

payoff_date_int_cell = doc.tables[1].rows[1].cells[1]
payoff_date_int_cell.text = 'Accrued Unpaid Interest as of ' + str(formatted_payoff_date)

prin_cell.text = locale.currency(bal, grouping = True)
int_cell.text = locale.currency(interest, grouping = True)
late_fee_cell.text =  locale.currency(misc, grouping = True)
net_bal_cell.text =  locale.currency(net_bal, grouping = True)

proceeds_cell.text = locale.currency(proceeds, grouping = True)

subtotal_cell.text = locale.currency(subtotal, grouping = True)

repo_cell.text = locale.currency(repo, grouping = True)
impound_cell.text = locale.currency(impo, grouping = True)
storing_cell.text = locale.currency(storing, grouping = True)
preparing_cell.text =  locale.currency(cop_veh, grouping = True)
selling_cell.text  = locale.currency(cos_veh, grouping = True)
court_cost_cell.text = locale.currency(court_cost, grouping = True)
other_cost_cell.text = locale.currency(other_other_cost, grouping = True)
total_cost_cell.text = locale.currency(total_cost, grouping = True)

gap_cell.text = locale.currency(abs(uacg), grouping = True)

total_credit_cell.text = locale.currency(uacg, grouping = True)

surplus_cell.text = "(" + locale.currency(abs(balance_due), grouping = True) + ")"

for row in table1.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(10)

# Adding the total underline portion for client
paragraph = doc.paragraphs[37]
run = paragraph.runs[6]
existing_text = run.text
run.text = existing_text[:-10]
added_text = locale.currency(abs(balance_due))
new_run = paragraph.add_run()
new_run.text = added_text
new_run.font.underline = True
new_run.font.size = Pt(10)

if abs(balance_due) != refund:
  print("Please review Defi Letter for account " + str(file_name))
  pass
else:
  doc.save(newdocx_path)
  print("Defi Letter for " + str(file_name) + " has been created")

